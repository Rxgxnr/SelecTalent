import streamlit as st
from openai import OpenAI
import fitz
import pandas as pd
from docx import Document
from io import BytesIO
import re

# --- Configuración Inicial ---
st.set_page_config(page_title="Bienvenido/a a SelecTalent", layout="centered")
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# --- Inicialización de Estado ---
for key in ["archivos_cv", "resultados", "descriptor", "nombre_cargo", "resumen_descriptor", "resumen"]:
    if key not in st.session_state:
        st.session_state[key] = [] if key in ["archivos_cv", "resultados", "resumen"] else ""

# --- Funciones ---
def extraer_texto_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "".join([page.get_text() for page in doc])
    except Exception as e:
        return f"❌ Error al leer PDF: {e}"

def generar_descriptor(p1, p2, p3):
    prompt = f"""Actúa como un Agente de Recursos Humanos experto.
Crea un Descriptor de Cargo profesional y detallado con base en:
1. Tipo de cargo: {p1}
2. Habilidades técnicas necesarias: {p2}
3. Perfil humano deseable: {p3}

Solo incluye texto estructurado (sin evaluaciones ni ranking)."""
    response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return response.choices[0].message.content.strip()

def generar_resumen_descriptor(descriptor):
    prompt = f"""Actúa como un especialista en Recursos Humanos.
Resume este Descriptor de Cargo de manera profesional y clara:
{descriptor}

No incluyas evaluaciones ni escalas numéricas."""
    response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return response.choices[0].message.content.strip()

def analizar_cv(descriptor, texto_cv):
    prompt = f"""Actúa como un Agente de Recursos Humanos experto.
Analiza el siguiente CV en función del descriptor entregado.

Descriptor del cargo:
{descriptor}

Currículum del candidato:
{texto_cv}

Indica:
- Evaluación académica
- Experiencia laboral
- Habilidades técnicas
- Competencias blandas
- Fortalezas y debilidades
- Recomendación general (avanzar o no a entrevista)

NO incluyas notas ni rankings de afinidad."""
    response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": prompt}])
    return response.choices[0].message.content.strip()

def generar_word(resultados, nombre_cargo):
    doc = Document()
    doc.add_heading(f"Reporte de Análisis - {nombre_cargo}", level=1)
    for r in resultados:
        doc.add_heading(r["nombre"], level=2)
        doc.add_paragraph(r["resultado"])
        doc.add_page_break()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, f"Reporte de Análisis - {nombre_cargo}.docx"

# --- Botón para reiniciar ---
if st.button("🔄 Consultar Otro Cargo"):
    st.session_state.clear()
    st.rerun()

# --- Título Principal ---
st.title("🤖 SelecTalent: Análisis de CV con IA")

# --- Mostrar estado actual ---
if st.session_state.get('descriptor'):
    st.sidebar.success(f"Cargo actual: {st.session_state.nombre_cargo}")

# --- Entrada de Descriptor ---
modo = st.radio("¿Quieres cargar un descriptor o prefieres que te ayude?", ["📂 Cargar Descriptor", "💬 Hacer Preguntas"])

if modo == "📂 Cargar Descriptor":
    archivo = st.file_uploader("Sube un descriptor en .txt o .pdf", type=["txt", "pdf"])
    if archivo is not None:
        try:
            descriptor = archivo.read().decode("utf-8") if archivo.type == "text/plain" else extraer_texto_pdf(archivo)
            st.session_state.descriptor = descriptor
            st.session_state.nombre_cargo = archivo.name.rsplit(".", 1)[0]

            with st.spinner("Generando resumen del descriptor..."):
                resumen = generar_resumen_descriptor(descriptor)

            st.session_state.resumen_descriptor = resumen
            st.success("✅ Descriptor cargado correctamente.")
            st.rerun()
        except Exception as e:
            st.error(f"Error al procesar el archivo: {str(e)}")

elif modo == "💬 Hacer Preguntas":
    with st.form("formulario"):
        p1 = st.text_input("¿Qué tipo de cargo buscas?")
        p2 = st.text_input("¿Qué habilidades o conocimientos debe tener?")
        p3 = st.text_input("¿Qué perfil humano o experiencia previa es deseable?")
        enviar = st.form_submit_button("Generar descriptor")
    if enviar:
        st.session_state.descriptor = generar_descriptor(p1, p2, p3)
        st.session_state.nombre_cargo = p1
        st.session_state.resumen_descriptor = generar_resumen_descriptor(st.session_state.descriptor)
        st.success("✅ Descriptor generado correctamente.")
        st.rerun()

# --- Análisis de CVs ---
if st.session_state.get('descriptor') and st.session_state.get('resumen_descriptor'):
    st.subheader(f"📝 Descriptor: {st.session_state.nombre_cargo}")
    with st.expander("Ver descriptor completo"):
        st.text_area("", st.session_state.descriptor, height=150, label_visibility="collapsed")
    st.info(f"📌 **Resumen del Descriptor:**\n{st.session_state.resumen_descriptor}")

    st.divider()
    st.subheader("📄 Carga los CVs en PDF")
    archivos_cv = st.file_uploader("Selecciona uno o varios archivos", type=["pdf"], accept_multiple_files=True)
    
if archivos_cv:
    if st.button("🔍 Analizar CVs"):
        st.session_state.archivos_cv = archivos_cv
        resultados = []
        for archivo in archivos_cv:
            texto = extraer_texto_pdf(archivo)
            if texto.startswith("❌"):
                st.error(f"{archivo.name}: {texto}")
                continue
            with st.spinner(f"Analizando {archivo.name}..."):
                analisis = analizar_cv(st.session_state.descriptor, texto)
                resultados.append({
                    "nombre": archivo.name,
                    "resultado": analisis
                })
            st.success(f"✅ CV '{archivo.name}' analizado con éxito")
        st.session_state.resultados = resultados
        st.experimental_rerun()

# --- Exportación ---
if st.session_state.get('resultados'):
    st.divider()
    st.subheader("📥 Exportar Resultados")
    col1, col2 = st.columns(2)

    with col1:
        df = pd.DataFrame([
            {
                "Nombre CV": r["nombre"],
                "Cargo": st.session_state.nombre_cargo,
            } for r in st.session_state.resultados
        ])
        excel_buffer = BytesIO()
        df.to_excel(excel_buffer, index=False)
        excel_buffer.seek(0)
        st.download_button("📊 Descargar Excel", excel_buffer, file_name=f"Resumen - {st.session_state.nombre_cargo}.xlsx")

    with col2:
        word_data, word_name = generar_word(st.session_state.resultados, st.session_state.nombre_cargo)
        st.download_button("📄 Descargar Word", word_data, file_name=word_name)
